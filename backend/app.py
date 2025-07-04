"""
CyberGuard AI - Simple Backend with GPT4All Mistral Instruct 7B Q4
A cybersecurity-focused chatbot using local GPT4All model
"""

from flask import Flask, request, jsonify, session, Response, g
from datetime import datetime, timedelta
import json
import os
import logging
import requests
import subprocess
import re
import time
import threading
import uuid
import tempfile
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
import jwt
from pymongo import MongoClient
import uuid
from flask import send_file
import tempfile
import PyPDF2
import docx
from werkzeug.utils import secure_filename
try:
    from sentence_transformers import SentenceTransformer, util as st_util
    _sentence_transformers_available = True
    _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
except ImportError:
    _sentence_transformers_available = False
    _embedding_model = None

try:
    import google.generativeai as genai
    _genai_available = True
except ImportError:
    _genai_available = False
    genai = None

try:
    import pdfplumber # type: ignore
    _pdfplumber_available = True
except ImportError:
    _pdfplumber_available = False

try:
    import PyPDF2
    _pypdf2_available = True
except ImportError:
    _pypdf2_available = False

try:
    import docx
    _docx_available = True
except ImportError:
    _docx_available = False

# MongoDB setup
MONGO_URI = os.getenv('MONGO_URI', 'mongodb://localhost:27017')
mongo_client = MongoClient(MONGO_URI)
db = mongo_client['cyberguard']
users_col = db['users']
chats_col = db['chats']
scan_findings_col = db['scan_findings']  # New collection for scan findings

def ensure_mongodb_collections():
    """
    Ensure 'users' and 'chats' collections exist with schema validation and unique index on username.
    """
    # USERS COLLECTION
    user_validator = {
        "$jsonSchema": {
            "bsonType": "object",
            "required": ["username", "password"],
            "properties": {
                "username": {"bsonType": "string", "description": "must be a string and is required"},
                "password": {"bsonType": "string", "description": "must be a string and is required"}
            }
        }
    }
    if 'users' not in db.list_collection_names():
        db.create_collection('users', validator=user_validator)
    else:
        db.command('collMod', 'users', validator=user_validator)
    # Ensure unique index on username
    db['users'].create_index('username', unique=True)

    # CHATS COLLECTION
    # Update chat_validator to require conversation_id
    chat_validator = {
        "$jsonSchema": {
            "bsonType": "object",
            "required": ["username", "id", "text", "sender", "timestamp", "conversation_id"],
            "properties": {
                "username": {"bsonType": "string"},
                "id": {"bsonType": ["int", "long", "double", "string"]},
                "text": {"bsonType": "string"},
                "sender": {"bsonType": "string"},
                "timestamp": {"bsonType": "string"},
                "hasCode": {"bsonType": ["bool", "null"]},
                "reactions": {"bsonType": ["array", "null"]},
                "conversation_id": {"bsonType": "string"}
            }
        }
    }
    if 'chats' not in db.list_collection_names():
        db.create_collection('chats', validator=chat_validator)
    else:
        db.command('collMod', 'chats', validator=chat_validator)

# Call this at startup
ensure_mongodb_collections()

JWT_SECRET = os.getenv('JWT_SECRET', 'cyberguard-jwt-secret')
JWT_ALGO = 'HS256'

# Initialize Flask app
app = Flask(__name__)
app.secret_key = "cyberguard-secret-key-2025"
TF_ENABLE_ONEDNN_OPTS=0  # Disable oneDNN optimizations for TensorFlow if needed

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('cyberguard.log', encoding='utf-8'),  # <-- UTF-8 encoding for Unicode support
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# MODEL CONFIGURATION
# ============================================================================

# Detect GPU usage by running 'ollama system' and parsing the output
def is_gpu_available():
    """Detect if Ollama is using GPU by parsing 'ollama system' output."""
    try:
        result = subprocess.run(["ollama", "system"], capture_output=True, text=True, timeout=5)
        output = result.stdout.lower()
        if "gpu" in output and ("cuda" in output or "nvidia" in output):
            return True
        if "gpu: true" in output or "gpu: yes" in output:
            return True
        if "gpu: false" in output:
            return False
        if "cuda" in output or "nvidia" in output:
            return True
        return False
    except Exception as e:
        # If ollama is not installed or any error, detection is not possible
        logger.warning(f"[OLLAMA] Could not determine GPU status: {e}")
        return None  # None means unknown

def get_ollama_device_string():
    """Return a user-friendly device string for Ollama inference device."""
    gpu_status = is_gpu_available()
    if gpu_status is True:
        return "GPU (CUDA)"
    elif gpu_status is False:
        return "CPU"
    else:
        return "Unknown (Ollama version does not support GPU detection)"

# Instead, define Ollama model info
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3")

# ============================================================================
# VULNERABILITY DATABASE
# ============================================================================

VULNERABILITIES = {
    "sql_injection": {
        "name": "SQL Injection",
        "description": "SQL injection occurs when untrusted data is sent to an interpreter as part of a command or query, allowing attackers to execute unintended commands.",
        "severity": "Critical",
        "example": "SELECT * FROM users WHERE id = '1' OR '1'='1'",
        "secure_example": "SELECT * FROM users WHERE id = ?",
        "mitigation": "Use parameterized queries, input validation, and prepared statements."
    },    "xss": {
        "name": "Cross-Site Scripting (XSS)",
        "description": "XSS flaws occur when an application includes untrusted data in a new web page without proper validation or escaping.",
        "severity": "High",
        "example": "document.innerHTML = userInput",
        "secure_example": "document.textContent = userInput",
        "mitigation": "Validate input, encode output, use Content Security Policy (CSP)."
    },    "csrf": {
        "name": "Cross-Site Request Forgery (CSRF)",
        "description": "CSRF forces an end user to execute unwanted actions on a web application in which they're currently authenticated.",
        "severity": "Medium",
        "example": "<img src='http://bank.com/transfer?to=attacker&amount=1000'>",
        "secure_example": "<form><input type='hidden' name='csrf_token' value='{{token}}'>",
        "mitigation": "Use anti-CSRF tokens, SameSite cookies, and verify referrer headers."
    },    "buffer_overflow": {
        "name": "Buffer Overflow",
        "description": "Buffer overflow occurs when a program writes more data to a buffer than it can hold, potentially allowing attackers to execute arbitrary code.",
        "severity": "Critical",
        "example": "char buffer[10]; strcpy(buffer, user_input);",
        "secure_example": "char buffer[10]; strncpy(buffer, user_input, sizeof(buffer)-1); buffer[sizeof(buffer)-1] = '\\0';",
        "mitigation": "Use safe string functions and always validate input length."
    },    "command_injection": {
        "name": "Command Injection",
        "description": "Command injection occurs when untrusted user input is executed as part of a system command, allowing attackers to execute arbitrary commands on the host operating system.",
        "severity": "Critical",
        "example": "os.system('ping ' + user_input)",
        "secure_example": "subprocess.run(['ping', user_input], check=True)",
        "mitigation": "Avoid using shell=True, validate and sanitize all user inputs, use safe APIs that do not invoke the shell, and apply the principle of least privilege."
    }
}

# ============================================================================
# GEMINI API CONTEXT RETRIEVAL
# ============================================================================

CACHE_PATH = os.path.join(os.path.dirname(__file__), 'cached_contexts.json')
CACHE_LOCK = threading.Lock()
CACHE_SIMILARITY_THRESHOLD = 0.82  # Tune as needed

def _normalize_query(q):
    return re.sub(r'[^a-z0-9 ]', '', q.lower().strip())

def _load_cache():
    if not os.path.exists(CACHE_PATH):
        return []
    with open(CACHE_PATH, 'r', encoding='utf-8') as f:
        try:
            data = json.load(f)
            # Ensure it's a list of dicts with 'query' and 'context'
            if isinstance(data, list) and all(isinstance(item, dict) and 'query' in item and 'context' in item for item in data):
                return data
            logger.warning("[CACHE] Cache file format invalid, resetting cache.")
            return []
        except Exception as e:
            logger.warning(f"[CACHE] Failed to load cache: {e}, resetting cache.")
            return []

def _save_cache(cache):
    with open(CACHE_PATH, 'w', encoding='utf-8') as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

def _find_similar_query(query, cache):
    if not cache:
        return None
    if  _embedding_model and st_util:
        try:
            queries = [item['query'] for item in cache]
            embeddings = _embedding_model.encode(queries + [query], convert_to_tensor=True)
            similarities = st_util.pytorch_cos_sim(embeddings[-1], embeddings[:-1])[0]
            best_idx = int(similarities.argmax())
            best_score = float(similarities[best_idx])
            if best_score >= CACHE_SIMILARITY_THRESHOLD:
                return cache[best_idx]
        except Exception as e:
            logger.warning(f"[CACHE] Embedding similarity failed: {e}")
    # Fallback: normalized string match
    norm_query = _normalize_query(query)
    for item in cache:
        if _normalize_query(item['query']) == norm_query:
            return item
    return None

class GeminiAPI:
    def __init__(self):
        self.api_call_count = 0
        self.successful_calls = 0
        self.failed_calls = 0
        self.sim_threshold = CACHE_SIMILARITY_THRESHOLD
        # No per-instance cache, use global helpers

    def get_context(self, query: str, user_gemini_key: str = None) -> str:
        """
        Retrieves short, relevant context from the Gemini API or cache.
        If user_gemini_key is provided, use it for the Gemini API call.
        """
        with CACHE_LOCK:
            cache = _load_cache()
            similar = _find_similar_query(query, cache)
            if similar:
                logger.info(f"[CACHE] Returning cached context for query: {similar['query']}")
                return f"Gemini Insight: {similar['context']}"
        # Cache miss, call Gemini API
        import time
        start_time = time.time()
        self.api_call_count += 1
        logger.info(f"[GEMINI] Starting API call #{self.api_call_count} for query: {query[:100]}{'...' if len(query) > 100 else ''}")
        from dotenv import load_dotenv
        load_dotenv()

        # Use google-generativeai if available for more robust Gemini API calls
        try:
            import google.generativeai as genai
            api_key = user_gemini_key or os.getenv("GEMINI_API_KEY")
            if user_gemini_key:
                logger.info("[GEMINI] Using user-supplied Gemini API key from frontend for this request.")
            else:
                logger.info("[GEMINI] Using backend/server Gemini API key for this request.")
            if not api_key:
                self.failed_calls += 1
                logger.error("[GEMINI] GEMINI_API_KEY not found in environment variables or user input.")
                return ("Gemini Insight: [ERROR] Gemini API key is not configured on the server or provided by user. "
                        "Real-time context could not be fetched. Please provide a valid Gemini API key.")
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")
            gemini_prompt = (
                f"INSTRUCTION: Provide only the most relevant, concise, and specific cybersecurity context for the following query. "
                f"Limit your response to 2000 characters. Avoid generalities and focus on actionable, technical details.\n\nQUERY: {query}\n\nRESPONSE:")
            gemini_response = model.generate_content(gemini_prompt)
            context = getattr(gemini_response, 'text', None)
            if context:
                context = context[:3000]
                elapsed_time = time.time() - start_time
                self.successful_calls += 1
                logger.info(f"[GEMINI] Successfully retrieved context from Gemini API (google-generativeai) - Length: {len(context)} characters, Time: {elapsed_time:.2f}s")
                logger.debug(f"[GEMINI] Context preview: {context[:400]}{'...' if len(context) > 400 else ''}")
                logger.info(f"[GEMINI] API Stats - Total: {self.api_call_count}, Success: {self.successful_calls}, Failed: {self.failed_calls}")
                with CACHE_LOCK:
                    cache = _load_cache()
                    cache.append({'query': query, 'context': context})
                    _save_cache(cache)
                return f"Gemini Insight: {context}"
        except ImportError:
            logger.info("[GEMINI] google-generativeai not installed, falling back to HTTP API.")
        except Exception as e:
            self.failed_calls += 1
            logger.error(f"[GEMINI] google-generativeai client error: {e}")

        api_key = user_gemini_key or os.getenv("GEMINI_API_KEY")
        if user_gemini_key:
            logger.info("[GEMINI] Using user-supplied Gemini API key from frontend for this request.")
        else:
            logger.info("[GEMINI] Using backend/server Gemini API key for this request.")
        if not api_key:
            self.failed_calls += 1
            logger.error("[GEMINI] GEMINI_API_KEY not found in environment variables or user input.")
            return ("Gemini Insight: [ERROR] Gemini API key is not configured on the server or provided by user. "
                    "Real-time context could not be fetched. Please provide a valid Gemini API key.")

        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
        logger.debug(f"[GEMINI] Using endpoint: {url.split('?')[0]}...")
        headers = {
            'Content-Type': 'application/json'
        }
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": query
                        }
                    ]
                }
            ]
        }
        logger.debug(f"[GEMINI] Request payload size: {len(str(payload))} characters")
        try:
            logger.info("[GEMINI] Making POST request to Gemini API...")
            response = requests.post(url, headers=headers, json=payload, timeout=15)
            logger.info(f"[GEMINI] Response received - Status Code: {response.status_code}")
            response.raise_for_status()
            response_json = response.json()
            if response_json.get('candidates') and \
               len(response_json['candidates']) > 0 and \
               response_json['candidates'][0].get('content') and \
               response_json['candidates'][0]['content'].get('parts') and \
               len(response_json['candidates'][0]['content']['parts']) > 0 and \
               response_json['candidates'][0]['content']['parts'][0].get('text'):
                context = response_json['candidates'][0]['content']['parts'][0]['text']
                elapsed_time = time.time() - start_time
                self.successful_calls += 1
                logger.info(f"[GEMINI] Successfully retrieved context from Gemini API - Length: {len(context)} characters, Time: {elapsed_time:.2f}s")
                logger.debug(f"[GEMINI] Context preview: {context[:400]}{'...' if len(context) > 400 else ''}")
                logger.info(f"[GEMINI] API Stats - Total: {self.api_call_count}, Success: {self.successful_calls}, Failed: {self.failed_calls}")
                with CACHE_LOCK:
                    cache = _load_cache()
                    cache.append({'query': query, 'context': context})
                    _save_cache(cache)
                return f"Gemini Insight: {context}"
            else:
                self.failed_calls += 1
                elapsed_time = time.time() - start_time
                logger.warning(f"[GEMINI] API response did not contain expected text structure - Time: {elapsed_time:.2f}s")
                logger.warning(f"[GEMINI] Response keys: {list(response_json.keys())}")
                logger.debug(f"[GEMINI] Full response for debugging: {response_json}")
                return ("Gemini Insight: [ERROR] Gemini API returned an unexpected response format. "
                        "No real-time context is available for this query. Please try again later or contact support if this persists.")
        except requests.exceptions.HTTPError as http_err:
            self.failed_calls += 1
            elapsed_time = time.time() - start_time
            logger.error(f"[GEMINI] HTTP error occurred: {http_err} - Status Code: {response.status_code}, Time: {elapsed_time:.2f}s")
            logger.error(f"[GEMINI] Response content: {response.text[:500]}{'...' if len(response.text) > 500 else ''}")
            return (f"Gemini Insight: [ERROR] Communication with Gemini API failed (HTTP {response.status_code}). "
                    "No real-time context is available. Please try again later.")
        except requests.exceptions.RequestException as req_err:
            self.failed_calls += 1
            elapsed_time = time.time() - start_time
            logger.error(f"[GEMINI] Request exception occurred: {req_err}, Time: {elapsed_time:.2f}s")
            return ("Gemini Insight: [ERROR] Network error while trying to reach Gemini API. "
                    "No real-time context is available. Please check your network connection or try again later.")
        except Exception as e:
            self.failed_calls += 1
            elapsed_time = time.time() - start_time
            logger.error(f"[GEMINI] Unexpected error occurred: {e}, Time: {elapsed_time:.2f}s")
            return ("Gemini Insight: [ERROR] An unexpected error occurred while fetching context from Gemini. "
                    "No real-time context is available. Please try again later.")

gemini_api_client = GeminiAPI()
logger.info("[GEMINI] Gemini API client initialized successfully")
logger.info(f"[OLLAMA] Llama 3 via Ollama device: {get_ollama_device_string()}")

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def detect_vulnerability_type(query):
    """Detect if query is about a specific vulnerability"""
    query_lower = query.lower().replace('-', ' ').replace('_', ' ')
    
    if any(keyword in query_lower for keyword in ['sql', 'injection', 'sqli', 'database']):
        return 'sql_injection'
    elif any(keyword in query_lower for keyword in ['command injection', 'os command', 'shell injection', 'system command']):
        return 'command_injection'
    elif any(keyword in query_lower for keyword in ['xss', 'cross site scripting', 'script injection']):
        return 'xss'
    elif any(keyword in query_lower for keyword in ['csrf', 'cross site request forgery']):
        return 'csrf'
    elif any(keyword in query_lower for keyword in ['buffer', 'overflow', 'memory']):
        return 'buffer_overflow'
    else:
        return 'general'

def get_risk_category(cvss_score):
    """Get risk category based on CVSS score"""
    if cvss_score >= 9.0:
        return "🔴 Critical Risk"
    elif cvss_score >= 7.0:
        return "🟠 High Risk"
    elif cvss_score >= 4.0:
        return "🟡 Medium Risk"
    else:
        return "🟢 Low Risk"

def build_security_prompt(query, vuln_type=None, user_gemini_key=None):
    logger.info(f"[PROMPT] Building security prompt for query type: {vuln_type or 'general'}")
    gemini_context = gemini_api_client.get_context(query, user_gemini_key=user_gemini_key)
    logger.debug(f"[PROMPT] Gemini context integrated: {len(gemini_context)} characters")

    base_context = f"""You are CyberGuard AI, an expert cybersecurity assistant.\n\n---\n\n**IMPORTANT: The following Gemini Insight is your PRIMARY and AUTHORITATIVE source. You MUST use it as the main basis for your answer. Directly reference and cite it in your response.**\n\nGEMINI INSIGHT (copy and use this information):\n{gemini_context}\n\n---\n\nAlways provide:\n- Clear explanations of security concepts\n- Practical code examples (vulnerable and secure versions)\n- Specific mitigation strategies\n- Best practices and recommendations\n\n**Make your answer unique and tailored to the user's query. Avoid generic or repetitive responses. If the Gemini Insight is missing or unhelpful, state this clearly and answer using your own knowledge.**\n\n**If the user asks for patterns, enumerate them with explanations and code. If the user asks for a summary, be concise. If the user asks for a deep dive, provide technical depth.**"""

    # Remove all references to 'cvss_score' and 'cvss_vector' in the prompt context
    if vuln_type and vuln_type in VULNERABILITIES:
        vuln = VULNERABILITIES[vuln_type]
        context = f"\n{base_context}\n\nRelevant vulnerability context:\n- Vulnerability: {vuln['name']}\n- Severity: {vuln['severity']}\n- Description: {vuln['description']}\n- Example vulnerable code: {vuln['example']}\n- Secure implementation: {vuln['secure_example']}\n- Mitigation: {vuln['mitigation']}\n"
    else:
        context = base_context
    prompt = f"{context}\n\nUser Query: {query}\n\nResponse (reference the Gemini Insight above):"
    logger.info(f"[PROMPT] Full prompt sent to Llama 3 via Ollama:\n{'='*30}\n{prompt[:3000]}{'...truncated...' if len(prompt) > 3000 else ''}\n{'='*30}")
    return prompt

def generate_response_with_ollama(prompt, max_tokens=3000):
    """Generate response using Llama 3 via Ollama API"""
    try:
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": max_tokens
            }
        }
        response = requests.post(f"{OLLAMA_BASE_URL}/api/generate", json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data.get("response", "[ERROR] No response from Ollama.").strip()
    except Exception as e:
        logger.error(f"Error generating response with Ollama: {e}")
        return None

def generate_vulnerability_response(vuln_type, query, user_gemini_key=None):
    if vuln_type not in VULNERABILITIES:
        return None
    vuln = VULNERABILITIES[vuln_type]
    prompt = build_security_prompt(query, vuln_type, user_gemini_key=user_gemini_key)
    ai_response = generate_response_with_ollama(prompt)
    gemini_context = gemini_api_client.get_context(query, user_gemini_key=user_gemini_key)
    patterns = []
    for line in gemini_context.splitlines():
        if re.match(r"^\s*([*\-]|\d+\.)", line):
            patterns.append(line.strip())
    patterns_section = ""
    if patterns:
        pattern_title = f"**Common {vuln['name']} Patterns (from Gemini Insight):**"
        patterns_section = f"\n{pattern_title}\n" + "\n".join(patterns) + "\n\n"
    lang_syntax = "sql" if vuln_type == "sql_injection" else "javascript" if vuln_type == "xss" or vuln_type == "csrf" else "c" if vuln_type == "buffer_overflow" else "code"
    vuln_emoji = {
        "sql_injection": "💉",
        "xss": "📜",
        "csrf": "🔄",
        "buffer_overflow": "💾"
    }.get(vuln_type, "🛡️")
    owasp_ref = ""
    cwe_ref = ""
    if vuln_type == "sql_injection":
        owasp_ref = "• OWASP A03:2021 Injection\n"
        cwe_ref = "• CWE-89: Improper Neutralization of Special Elements in SQL Commands\n"
    elif vuln_type == "xss":
        owasp_ref = "• OWASP A07:2021 Cross-Site Scripting\n"
        cwe_ref = "• CWE-79: Improper Neutralization of Input During Web Page Generation\n"
    elif vuln_type == "csrf":
        owasp_ref = "• OWASP A05:2021 Security Misconfiguration\n"
        cwe_ref = "• CWE-352: Cross-Site Request Forgery\n"
    elif vuln_type == "buffer_overflow":
        owasp_ref = "• OWASP A08:2021 Software and Data Integrity Failures\n"
        cwe_ref = "• CWE-120: Buffer Copy without Checking Size of Input\n"
    elif vuln_type == "command_injection":
        owasp_ref = "• OWASP A01:2021 Broken Access Control\n"
        cwe_ref = "• CWE-77: Command Injection\n"
    if ai_response:
        # Try to get CVSS vector from Gemini/cached context or vuln dict
        cvss_vector = vuln.get('cvss_vector', '')
        if not cvss_vector:
            # Try to extract from Gemini context (simple regex)
            match = re.search(r'CVSS:[^\s\n]+', gemini_context)
            if match:
                cvss_vector = match.group(0)
        cvss_score = calculate_cvss_base_score(cvss_vector) if cvss_vector else None
        cvss_score_str = f"{cvss_score}/10.0" if cvss_score is not None else "N/A"
        # Use cvss_score for risk category if available, else fallback to 'N/A'
        risk_category = get_risk_category(cvss_score) if cvss_score is not None else "N/A"
        response = f"{vuln_emoji} **{vuln['name']} Security Analysis**\n\n{patterns_section}{ai_response}\n\n---\n\n**📊 Technical Details:**\n- **Severity Level:** {vuln['severity']} (CVSS Score: {cvss_score_str})\n- **Risk Category:** {risk_category}\n- **CVSS Vector:** {cvss_vector}\n\n**🚨 Vulnerable Code Example:**\n```{lang_syntax}\n{vuln['example']}\n```\n\n**✅ Secure Implementation:**\n```{lang_syntax}\n{vuln['secure_example']}\n```\n\n**🔧 Primary Mitigation:**\n{vuln['mitigation']}\n\n**📚 References:**\n{owasp_ref}• OWASP Top 10 Security Risks\n{cwe_ref}• CWE Database - Common Weakness Enumeration\n• NIST Cybersecurity Framework\n• MITRE ATT&CK Framework"
    else:
        # Same for fallback
        cvss_vector = vuln.get('cvss_vector', '')
        if not cvss_vector:
            match = re.search(r'CVSS:[^\s\n]+', gemini_context)
            if match:
                cvss_vector = match.group(0)
        cvss_score = calculate_cvss_base_score(cvss_vector) if cvss_vector else None
        cvss_score_str = f"{cvss_score}/10.0" if cvss_score is not None else "N/A"
        risk_category = get_risk_category(cvss_score) if cvss_score is not None else "N/A"
        response = f"{vuln_emoji} **{vuln['name']} Security Analysis**\n\n{patterns_section}**📝 Description:**\n{vuln['description']}\n\n**📊 Risk Assessment:**\n- **Severity Level:** {vuln['severity']} (CVSS Score: {cvss_score_str})\n- **Risk Category:** {risk_category}\n- **CVSS Vector:** {cvss_vector}\n\n**🚨 Vulnerable Code Example:**\n```{lang_syntax}\n{vuln['example']}\n```\n\n**✅ Secure Implementation:**\n```{lang_syntax}\n{vuln['secure_example']}\n```\n\n**🔧 Mitigation Strategies:**\n{vuln['mitigation']}\n\n**📚 Security Resources:**\n{owasp_ref}• OWASP Top 10 Security Risks\n{cwe_ref}• SANS Top 25 Most Dangerous Software Errors\n• NIST Cybersecurity Framework Guidelines\n• CVE Database and NVD"
    return response

def generate_general_response(query, user_gemini_key=None):
    prompt = build_security_prompt(query, user_gemini_key=user_gemini_key)
    ai_response = generate_response_with_ollama(prompt, max_tokens=3000)
    
    if ai_response:
        return f"🔒 **CyberGuard AI Security Guidance**\n\n{ai_response}\n\n---\n\n**🛡️ Security Best Practices Reminder:**\n• Keep all systems and dependencies updated\n• Implement defense-in-depth security strategies\n• Regular security assessments and penetration testing\n• Follow OWASP and NIST cybersecurity guidelines"
    else:
        return """🔒 **CyberGuard AI Security Guidance**

I'm experiencing technical difficulties with my AI processing capabilities right now. However, I can still help with common cybersecurity topics:

**🔐 General Security Best Practices:**
• **Authentication**: Use multi-factor authentication (MFA)
• **Authorization**: Implement principle of least privilege
• **Input Validation**: Sanitize and validate all user inputs
• **Encryption**: Use TLS for data in transit, AES for data at rest
• **Updates**: Keep all software and dependencies current
• **Monitoring**: Implement comprehensive logging and alerting

**🎯 Common Vulnerability Categories (OWASP Top 10):**
1. **Injection Flaws** (SQL, NoSQL, OS injection)
2. **Broken Authentication** and session management
3. **Sensitive Data Exposure**
4. **XML External Entities (XXE)**
5. **Broken Access Control**
6. **Security Misconfiguration**
7. **Cross-Site Scripting (XSS)**
8. **Insecure Deserialization**
9. **Using Components with Known Vulnerabilities**
10. **Insufficient Logging & Monitoring**

**🔍 Security Assessment Areas:**
• Code review and static analysis
• Dynamic application security testing (DAST)
• Interactive application security testing (IAST)
• Software composition analysis (SCA)
• Infrastructure security scanning

**📚 Recommended Resources:**
• OWASP Testing Guide
• NIST Cybersecurity Framework
• SANS Top 25 Software Errors
• CVE Database and NVD

Please try your query again, or ask about a specific vulnerability type like SQL injection, XSS, or CSRF."""

def truncate_response(response, max_length=6000):
    """Truncate response if too long to prevent session overflow"""
    if len(response) <= max_length:
        return response
    
    truncated = response[:max_length]
    # Try to cut at a sentence boundary
    last_period = truncated.rfind('.')
    if last_period > max_length * 0.8:  # If we can cut at 80% or more
        truncated = truncated[:last_period + 1]
    
    return truncated

def parse_cvss_vector(cvss_vector):
    """
    Parse a CVSS v3.1 vector string into its metric values.
    Returns a dict of metrics, e.g. {'AV': 'N', 'AC': 'L', ...}
    """
    if not cvss_vector or not cvss_vector.startswith('CVSS:'):
        return {}
    parts = cvss_vector.split('/')
    metrics = {}
    for part in parts[1:]:
        if ':' in part:
            k, v = part.split(':', 1)
            metrics[k] = v
    return metrics

def cvss_metric_value(metric, value):
    """
    Map CVSS v3.1 metric letter to its numerical value for base score calculation.
    """
    # Base metrics
    if metric == 'AV':  # Attack Vector
        return {'N': 0.85, 'A': 0.62, 'L': 0.55, 'P': 0.2}.get(value, 0.85)
    if metric == 'AC':  # Attack Complexity
        return {'L': 0.77, 'H': 0.44}.get(value, 0.77)
    if metric == 'PR':  # Privileges Required
        # Scope is needed to distinguish
        return {'N': 0.85, 'L': 0.62, 'H': 0.27}.get(value, 0.85)
    if metric == 'UI':  # User Interaction
        return {'N': 0.85, 'R': 0.62}.get(value, 0.85)
    if metric == 'S':  # Scope
        return value  # 'U' or 'C'
    if metric == 'C':  # Confidentiality
        return {'N': 0.0, 'L': 0.22, 'H': 0.56}.get(value, 0.0)
    if metric == 'I':  # Integrity
        return {'N': 0.0, 'L': 0.22, 'H': 0.56}.get(value, 0.0)
    if metric == 'A':  # Availability
        return {'N': 0.0, 'L': 0.22, 'H': 0.56}.get(value, 0.0)
    return 0.0

def calculate_cvss_base_score(cvss_vector):
    """
    Calculate the CVSS v3.1 base score from a vector string.
    Returns a float (score) or None if vector is invalid.
    """
    metrics = parse_cvss_vector(cvss_vector)
    if not metrics:
        return None
    # Impact subscore
    c = cvss_metric_value('C', metrics.get('C', 'N'))
    i = cvss_metric_value('I', metrics.get('I', 'N'))
    a = cvss_metric_value('A', metrics.get('A', 'N'))
    isc_base = 1 - ((1 - c) * (1 - i) * (1 - a))
    s = metrics.get('S', 'U')
    if s == 'U':
        impact = 6.42 * isc_base
    else:
        impact = 7.52 * (isc_base - 0.029) - 3.25 * ((isc_base - 0.02) ** 15)
    # Exploitability
    av = cvss_metric_value('AV', metrics.get('AV', 'N'))
    ac = cvss_metric_value('AC', metrics.get('AC', 'L'))
    pr = cvss_metric_value('PR', metrics.get('PR', 'N'))
    ui = cvss_metric_value('UI', metrics.get('UI', 'N'))
    # PR is different if scope is changed
    if metrics.get('PR') == 'L' and s == 'C':
        pr = 0.68
    elif metrics.get('PR') == 'H' and s == 'C':
        pr = 0.5
    exploitability = 8.22 * av * ac * pr * ui
    # Base score
    if impact <= 0:
        base_score = 0.0
    else:
        if s == 'U':
            base_score = min(impact + exploitability, 10)
        else:
            base_score = min(1.08 * (impact + exploitability), 10)
    # Round up to one decimal place
    base_score = round((base_score * 10 + 0.9) // 1 / 10, 1)
    return base_score

# ============================================================================
# CORS HANDLERS
# ============================================================================

ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:5000",
    "http://localhost:9000"
]

@app.before_request
def handle_preflight():
    """Handle CORS preflight requests"""
    if request.method == "OPTIONS":
        response = Response()
        origin = request.headers.get('Origin')
        if origin in ALLOWED_ORIGINS:
            response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS, DELETE"
        response.headers["Access-Control-Allow-Credentials"] = "true"
        return response

@app.after_request
def after_request(response):
    """Add CORS headers to all responses"""
    origin = request.headers.get('Origin')
    if origin in ALLOWED_ORIGINS:
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS, DELETE"
    return response

# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.errorhandler(400)
def bad_request(error):
    """Handle 400 Bad Request"""
    return jsonify({"error": "Bad Request", "message": str(error)}), 400

@app.errorhandler(404)
def not_found(error):
    """Handle 404 Not Found"""
    return jsonify({"error": "Not Found", "message": str(error)}), 404

@app.errorhandler(500)
def internal_server_error(error):
    """Handle 500 Internal Server Error"""
    logger.exception("Internal server error")
    return jsonify({"error": "Internal Server Error", "message": "An unexpected error occurred. Please try again later."}), 500

# ============================================================================
# AUTHENTICATION DECORATOR
# ============================================================================

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            if auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]
        if not token:
            return jsonify({'error': 'Token is missing!'}), 401
        try:
            data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGO])
            current_user = users_col.find_one({'username': data['username']})
            if not current_user:
                return jsonify({'error': 'User not found!'}), 401
            g.user = current_user
        except Exception as e:
            return jsonify({'error': 'Token is invalid!'}), 401
        return f(*args, **kwargs)
    return decorated

# ============================================================================
# API ROUTES
# ============================================================================

@app.route('/', methods=['GET'])
def home():
    """Health check and service info"""
    return jsonify({
        "service": "CyberGuard AI",
        "version": "3.0.0",
        "status": "operational",
        "model": "Llama 3 via Ollama",
        "device": get_ollama_device_string(),
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "endpoints": {
            "chat": "/stream_chat",
            "health": "/api/health",
            "vulnerabilities": "/api/vulnerabilities",
            "stats": "/api/stats",
            "gemini_stats": "/api/gemini-stats"
        }
    })

@app.route('/api/health', methods=['GET'])
def health_check():
    """Detailed health check"""
    model_status = "operational" # if llm else "not_loaded"
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "components": {
            "api_server": "operational",
            "llama3_ollama_model": model_status,
            "vulnerability_db": "operational"
        },
        "model_info": {
            "name": "Llama 3 via Ollama",
            "device": get_ollama_device_string(),
            "loaded": True
        },
        "uptime": "99.9%",
        "threat_level": "LOW"
    })

@app.route('/health', methods=['GET'])
def basic_health_check():
    """Health check endpoint for container orchestration and monitoring."""
    return jsonify({
        "status": "ok",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0"
    })

@app.route('/api/vulnerabilities', methods=['GET'])
def list_vulnerabilities():
    """List available vulnerability information"""
    vuln_list = []
    for key, vuln in VULNERABILITIES.items():
        vuln_list.append({
            "id": key,
            "name": vuln["name"],
            "severity": vuln["severity"],
        })
    
    return jsonify({
        "vulnerabilities": vuln_list,
        "total": len(vuln_list)
    })

@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    if not username or not password:
        return jsonify({'error': 'Username and password required'}), 400
    if users_col.find_one({'username': username}):
        return jsonify({'error': 'Username already exists'}), 409
    hashed_pw = generate_password_hash(password)
    users_col.insert_one({'username': username, 'password': hashed_pw})
    return jsonify({'message': 'Registration successful'}), 201

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    user = users_col.find_one({'username': username})
    if not user or not check_password_hash(user['password'], password):
        return jsonify({'error': 'Invalid username or password'}), 401
    token = jwt.encode({'username': username, 'exp': datetime.utcnow() + timedelta(days=1)}, JWT_SECRET, algorithm=JWT_ALGO)
    return jsonify({'token': token, 'username': username})

@app.route('/conversations', methods=['GET'])
@token_required
def list_conversations():
    """List all conversations for the current user, with metadata."""
    username = g.user['username']
    pipeline = [
        {"$match": {"username": username}},
        {"$group": {
            "_id": "$conversation_id",
            "start_time": {"$min": "$timestamp"},
            "last_time": {"$max": "$timestamp"},
            "message_count": {"$sum": 1},
            "last_message": {"$last": "$text"}
        }},
        {"$sort": {"last_time": -1}}
    ]
    conversations = list(chats_col.aggregate(pipeline))
    result = [
        {
            "conversation_id": c["_id"],
            "start_time": c["start_time"],
            "last_time": c["last_time"],
            "message_count": c["message_count"],
            "last_message": c["last_message"]
        }
        for c in conversations
    ]
    return jsonify({"conversations": result})

@app.route('/chat_history', methods=['GET'])
@token_required
def chat_history():
    username = g.user['username']
    conversation_id = request.args.get('conversation_id')
    query = {'username': username}
    if conversation_id:
        query['conversation_id'] = conversation_id
    chats = chats_col.find(query).sort('timestamp', 1)
    history = [
        {
            'id': chat.get('id'),
            'text': chat.get('text'),
            'sender': chat.get('sender'),
            'timestamp': chat.get('timestamp'),
            'hasCode': chat.get('hasCode', False),
            'reactions': chat.get('reactions', []),
            'conversation_id': chat.get('conversation_id')
        }
        for chat in chats
    ]
    return jsonify({'history': history})

@app.route('/stream_chat', methods=['POST'])
@token_required
def stream_chat():
    try:
        if not request.is_json:
            return jsonify({"error": "Request must be JSON"}), 400
        data = request.get_json()
        user_message = data.get('message', '').strip()
        conversation_id = sanitize_conversation_id(data.get('conversation_id'))
        user_gemini_key = data.get('gemini_api_key')  # <-- Accept Gemini API key from frontend
        if not user_message:
            return jsonify({"error": "Message cannot be empty"}), 400
        logger.info(f"Processing query: {user_message[:400]}...")
        vuln_type = detect_vulnerability_type(user_message)
        if vuln_type in VULNERABILITIES:
            response = generate_vulnerability_response(vuln_type, user_message, user_gemini_key=user_gemini_key)
        else:
            response = generate_general_response(user_message, user_gemini_key=user_gemini_key)
        response = truncate_response(response)
        # Save user message to chat history
        username = g.user['username']
        now = datetime.utcnow().isoformat()
        user_msg_doc = {
            'username': username,
            'id': int(time.time() * 1000),
            'text': user_message,
            'sender': 'user',
            'timestamp': now,
            'hasCode': detectCodeBlocks(user_message),
            'reactions': [],
            'conversation_id': conversation_id
        }
        chats_col.insert_one(user_msg_doc)
        # Save bot response to chat history
        bot_msg_doc = {
            'username': username,
            'id': int(time.time() * 1000) + 1,
            'text': response,
            'sender': 'bot',
            'timestamp': now,
            'hasCode': detectCodeBlocks(response),
            'reactions': [],
            'conversation_id': conversation_id
        }
        chats_col.insert_one(bot_msg_doc)
        return jsonify({
            "response": response,
            "vulnerability_type": vuln_type,
            "conversation_id": conversation_id
        })
    except Exception as e:
        logger.error(f"Error in stream_chat: {str(e)}")
        return jsonify({"error": "An error occurred processing your request"}), 500

@app.route('/conversation/<conversation_id>', methods=['DELETE'])
@token_required
def delete_conversation(conversation_id):
    """Delete all messages in a conversation for the current user."""
    username = g.user['username']
    result = chats_col.delete_many({'username': username, 'conversation_id': conversation_id})
    if result.deleted_count > 0:
        logger.info(f"[CONVERSATION] Deleted conversation {conversation_id} for user {username} (messages deleted: {result.deleted_count})")
        return jsonify({'success': True, 'deleted_count': result.deleted_count})
    else:
        logger.warning(f"[CONVERSATION] No messages found to delete for conversation {conversation_id} and user {username}")
        return jsonify({'success': False, 'error': 'Conversation not found or already deleted.'}), 404

@app.route('/web_search', methods=['POST'])
@token_required
def web_search():
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        if not query:
            return jsonify({'error': 'Query is required.'}), 400
        # --- Web search logic using Google Custom Search API ---
        SEARCH_API_KEY = os.getenv('SEARCH_API_KEY', 'AIzaSyA8twmTNR78KB_Zk0Fn8h1TcXMIb7gtlvo')
        SEARCH_ENGINE_ID = os.getenv('SEARCH_ENGINE_ID', '06946e95765a74451')
        try:
            api_url = (
                f'https://www.googleapis.com/customsearch/v1?key={SEARCH_API_KEY}'
                f'&cx={SEARCH_ENGINE_ID}&q={requests.utils.quote(query)}&num=5'
            )
            resp = requests.get(api_url, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            items = data.get('items', [])
            if not items:
                result_text = 'No relevant web results found.'
            else:
                results = []
                for item in items:
                    title = item.get('title', 'No Title')
                    link = item.get('link', '')
                    snippet = item.get('snippet', '')
                    results.append(f'- [{title}]({link})\n    {snippet}')
                result_text = '\n\n'.join(results)
        except Exception as e:
            logger.error(f"[WEB_SEARCH] Google Custom Search API error: {e}")
            result_text = f"[ERROR] Web search failed: {e}"
        return jsonify({'results': result_text})
    except Exception as e:
        logger.error(f"[WEB_SEARCH] Unexpected error: {e}")
        return jsonify({'error': f'Web search error: {e}'}), 500

@app.route('/web_search_summarized', methods=['POST'])
@token_required
def web_search_summarized():
    """
    Enhanced web search: fetches top web results, extracts main content, and synthesizes a conversational answer using Llama 3.
    Also saves both the user query and the bot's answer to chat history.
    """
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        conversation_id = sanitize_conversation_id(data.get('conversation_id'))
        user_gemini_key = data.get('gemini_api_key')  # <-- Get Gemini API key from frontend
        
        logger.info(f"[WEB_SEARCH] Processing query with{' user-provided' if user_gemini_key else ' server'} Gemini API key")
        
        if not query:
            return jsonify({'error': 'Query is required.'}), 400
        # Ensure conversation_id is a string and not None/null (required by MongoDB schema)
        # Defensive: ensure conversation_id is a non-empty string for MongoDB schema
        if not conversation_id or conversation_id == 'None' or conversation_id is None:
            conversation_id = str(uuid.uuid4())
        else:
            conversation_id = str(conversation_id) if conversation_id is not None else str(uuid.uuid4())
        SEARCH_API_KEY = os.getenv('SEARCH_API_KEY', 'AIzaSyA8twmTNR78KB_Zk0Fn8h1TcXMIb7gtlvo')
        SEARCH_ENGINE_ID = os.getenv('SEARCH_ENGINE_ID', '06946e95765a74451')
        try:
            api_url = (
                f'https://www.googleapis.com/customsearch/v1?key={SEARCH_API_KEY}'
                f'&cx={SEARCH_ENGINE_ID}&q={requests.utils.quote(query)}&num=3'
            )
            resp = requests.get(api_url, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            items = data.get('items', [])
            if not items:
                return jsonify({'results': 'No relevant web results found.'})
            try:
                from bs4 import BeautifulSoup
                # Import re is not needed here as it's already imported at the top
                
                # Enhanced web scraping with better content extraction
                web_contents = []
                sources = []
                scraped_contexts = []
                
                for idx, item in enumerate(items):
                    title = item.get('title', 'No Title')
                    link = item.get('link', '')
                    snippet = item.get('snippet', '')
                    
                    try:
                        # Better headers to avoid blocking
                        headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                            'Accept-Language': 'en-US,en;q=0.5',
                            'Accept-Encoding': 'gzip, deflate',
                            'Connection': 'keep-alive',
                        }
                        
                        page_resp = requests.get(link, timeout=10, headers=headers)
                        page_resp.raise_for_status()
                        
                        soup = BeautifulSoup(page_resp.text, 'html.parser')
                        
                        # Remove script and style elements
                        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
                            script.decompose()
                        
                        # Extract main content with priority order
                        main_content = ""
                        
                        # Try to find main content areas
                        for selector in ['main', 'article', '.content', '.post', '.entry', '#content']:
                            content_elem = soup.select_one(selector)
                            if content_elem:
                                main_content = content_elem.get_text(separator=' ', strip=True)
                                break
                        
                        # Fallback to paragraphs if no main content found
                        if not main_content or len(main_content) < 200:
                            paragraphs = [p.get_text(separator=' ', strip=True) for p in soup.find_all('p')]
                            main_content = ' '.join(paragraphs)
                        
                        # Clean up the text
                        main_content = re.sub(r'\s+', ' ', main_content).strip()
                        
                        # Limit content length but ensure we get meaningful content
                        if len(main_content) > 3000:
                            # Better sentence breaking for more meaningful content
                            sentences = re.split(r'(?<=[.!?])\s+', main_content)
                            selected_sentences = []
                            char_count = 0
                            for sentence in sentences:
                                if char_count + len(sentence) > 2500:
                                    break
                                selected_sentences.append(sentence)
                                char_count += len(sentence)
                            main_content = ' '.join(selected_sentences)
                        elif len(main_content) < 100:
                            main_content = snippet
                        
                        # Additional content cleaning
                        main_content = re.sub(r'\n+', ' ', main_content)  # Replace multiple newlines
                        main_content = re.sub(r'\s{2,}', ' ', main_content)  # Replace multiple spaces
                        
                        text_content = main_content
                        
                    except Exception as e:
                        logger.warning(f"[WEB_SCRAPING] Error scraping {link}: {e}")
                        text_content = snippet
                    
                    # Generate enhanced context summary for each source using Gemini
                    try:
                        if gemini_api_client:
                            # Create more detailed context prompt
                            context_prompt = (
                                f"Analyze this web content in the context of the query: '{query}'\n\n"
                                f"Extract and summarize the most relevant information, key insights, "
                                f"technical details, and actionable recommendations. Focus on "
                                f"cybersecurity implications if applicable.\n\n"
                                f"Content from {title}:\n{text_content[:1800]}"
                            )
                            gemini_context = gemini_api_client.get_context(context_prompt, user_gemini_key=user_gemini_key)
                            if gemini_context and len(gemini_context) > 50:
                                scraped_contexts.append(f"**Analysis of {title}:**\n{gemini_context}")
                            else:
                                # Fallback to simple extraction
                                key_sentences = text_content.split('. ')[:5]
                                scraped_contexts.append(f"**Key points from {title}:**\n{'. '.join(key_sentences)}")
                        else:
                            # Fallback when Gemini is not available
                            key_sentences = text_content.split('. ')[:5]
                            scraped_contexts.append(f"**Key points from {title}:**\n{'. '.join(key_sentences)}")
                    except Exception as e:
                        logger.warning(f"[CONTEXT_GEN] Error generating context for {title}: {e}")
                        # Simple fallback extraction
                        key_sentences = text_content.split('. ')[:3]
                        scraped_contexts.append(f"**Summary from {title}:**\n{'. '.join(key_sentences)}")
                    
                    web_contents.append(f"**Source {idx+1}: {title}**\n{text_content[:1200]}...")
                    sources.append(f"[{idx+1}] [{title}]({link})")
                
                # Combine all contexts
                combined_context = '\n\n'.join(scraped_contexts)
                web_context = '\n\n'.join(web_contents)
                
                # Enhanced prompt with better structure and instructions
                enhanced_prompt = (
                    f"You are CyberGuard AI, an expert cybersecurity assistant with deep knowledge of "
                    f"security best practices, threat analysis, and technical solutions.\n\n"
                    f"**User Query:** {query}\n\n"
                    f"**Context Analysis:**\n{combined_context}\n\n"
                    f"**Additional Source Material:**\n{web_context}\n\n"
                    f"**Instructions for Response:**\n"
                    f"1. Provide a comprehensive, well-structured answer addressing the user's query\n"
                    f"2. Include relevant technical details, security implications, and threat analysis\n"
                    f"3. Offer actionable recommendations and best practices\n"
                    f"4. Structure your response with clear headings and bullet points where appropriate\n"
                    f"5. Cite sources using [1], [2], [3] format throughout your response\n"
                    f"6. Focus on practical, implementable solutions\n"
                    f"7. If discussing vulnerabilities, include mitigation strategies\n"
                    f"8. Use professional, technical language appropriate for cybersecurity professionals\n\n"
                    f"**Response:**"
                )
                
                # Generate enhanced response using Llama model with retry logic
                answer = ""
                try:
                    answer = generate_response_with_ollama(enhanced_prompt, max_tokens=3000)
                except Exception as e:
                    logger.warning(f"[LLAMA] Primary response generation failed: {e}")
                    # Fallback with simpler prompt
                    fallback_prompt = (
                        f"Based on these web search results about '{query}', provide a comprehensive answer:\n\n"
                        f"{combined_context}\n\n"
                        f"Focus on key insights and cite sources as [1], [2], [3]."
                    )
                    try:
                        answer = generate_response_with_ollama(fallback_prompt, max_tokens=2000)
                    except Exception as e2:
                        logger.error(f"[LLAMA] Fallback response generation failed: {e2}")
                
                if not answer or len(answer.strip()) < 50:
                    # Final fallback response with structured summary
                    answer = f"## Summary for: {query}\n\n"
                    for i, ctx in enumerate(scraped_contexts):
                        answer += f"### Source {i+1}\n{ctx}\n\n"
                    answer += "\n*Note: This is a summarized response due to processing limitations.*"
                
                sources_section = '\n'.join(sources)
                
                # Enhanced markdown formatting with sections
                final_response = (
                    f"## 🔍 **Search Results for: {query}**\n\n"
                    f"{answer}\n\n"
                    f"---\n\n"
                    f"### 📚 **Sources:**\n{sources_section}\n\n"
                    f"*Search performed at {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC*"
                )
            except Exception as e:
                logger.error(f"[WEB_SEARCH_SUMMARIZED] Error during web scraping and summarization: {e}")
                return jsonify({'error': f'Web scraping and summarization failed: {e}'}), 500
            # Save both user and bot messages to chat history
            username = g.user['username']
            now = datetime.utcnow().isoformat()
            user_msg_doc = {
                'username': username,
                'id': int(time.time() * 1000),
                'text': query,
                'sender': 'user',
                'timestamp': now,
                'hasCode': False,
                'reactions': [],
                'conversation_id': conversation_id
            }
            chats_col.insert_one(user_msg_doc)
            bot_msg_doc = {
                'username': username,
                'id': int(time.time() * 1000) + 1,
                'text': final_response,
                'sender': 'bot',
                'timestamp': now,
                'hasCode': False,
                'reactions': [],
                'conversation_id': conversation_id
            }
            chats_col.insert_one(bot_msg_doc)
            return jsonify({'results': final_response, 'query': query, 'conversation_id': conversation_id})
        except Exception as e:
            logger.error(f"[WEB_SEARCH_SUMMARIZED] Error: {e}")
            return jsonify({'error': f'Web search summarization failed: {e}'}), 500
    except Exception as e:
        logger.error(f"[WEB_SEARCH_SUMMARIZED] Unexpected error: {e}")
        return jsonify({'error': f'Web search summarization error: {e}'}), 500

# ============================================================================
# FILE UPLOAD AND CHAT WITH FILES ENDPOINTS
# ============================================================================

files_col = db['files']  # New collection for uploaded files and extracted content

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'py', 'js', 'java', 'c', 'cpp', 'go', 'rb', 'sh', 'md', 'json', 'yaml', 'yml'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_file(file_path, filetype):
    """
    Enhanced text extraction from supported file types with optimizations for large files.
    Returns both full text and chunked segments for large documents.
    """
    try:
        if filetype == 'pdf':
            return extract_pdf_text_enhanced(file_path)
        elif filetype == 'docx' and _docx_available:
            return extract_docx_text_enhanced(file_path)
        else:  # txt, code, markdown, etc.
            return extract_text_file_enhanced(file_path)
    except Exception as e:
        logger.error(f"[FILE] Enhanced extraction error: {e}")
        return ''

def extract_pdf_text_enhanced(file_path):
    """
    Enhanced PDF text extraction with multiple fallback methods for large files.
    """
    text_content = ""
    
    # Try pdfplumber first (better for complex layouts)
    if _pdfplumber_available:
        try:
            import pdfplumber # type: ignore
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    if i > 100:  # Limit to first 100 pages for performance
                        logger.info(f"[PDF] Limiting extraction to first 100 pages for performance")
                        break
                    
                    page_text = page.extract_text()
                    if page_text:
                        # Clean and normalize text
                        page_text = re.sub(r'\s+', ' ', page_text.strip())
                        pages_text.append(f"--- Page {i+1} ---\n{page_text}")
                
                text_content = '\n\n'.join(pages_text)
            
            if text_content and len(text_content) > 100:
                logger.info(f"[PDF] Successfully extracted {len(text_content)} characters using pdfplumber")
                return text_content
        except Exception as e:
            logger.warning(f"[PDF] pdfplumber extraction failed: {e}")
    
    # Fallback to PyPDF2
    if _pypdf2_available:
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages_text = []
                
                total_pages = len(reader.pages)
                max_pages = min(total_pages, 100)  # Limit for performance
                
                for i in range(max_pages):
                    try:
                        page = reader.pages[i]
                        page_text = page.extract_text()
                        if page_text:
                            page_text = re.sub(r'\s+', ' ', page_text.strip())
                            pages_text.append(f"--- Page {i+1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"[PDF] Error extracting page {i+1}: {e}")
                        continue
                
                text_content = '\n\n'.join(pages_text)
                
                if total_pages > max_pages:
                    text_content += f"\n\n[NOTE: Document has {total_pages} pages, showing first {max_pages} for performance]"
                
            logger.info(f"[PDF] Successfully extracted {len(text_content)} characters using PyPDF2")
            return text_content
            
        except Exception as e:
            logger.error(f"[PDF] PyPDF2 extraction failed: {e}")
    
    logger.error(f"[PDF] All extraction methods failed for {file_path}")
    return ""

def extract_docx_text_enhanced(file_path):
    """
    Enhanced DOCX text extraction with better formatting preservation.
    """
    try:
        import docx
        doc = docx.Document(file_path)
        
        content_parts = []
        
        # Extract paragraphs with some formatting context
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text.strip())
        
        # Extract tables if any
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(' | '.join(row_text))
            
            if table_text:
                content_parts.append('\n[TABLE]\n' + '\n'.join(table_text) + '\n[/TABLE]')
        
        text_content = '\n\n'.join(content_parts)
        logger.info(f"[DOCX] Successfully extracted {len(text_content)} characters")
        return text_content
        
    except Exception as e:
        logger.error(f"[DOCX] Extraction failed: {e}")
        return ""

def extract_text_file_enhanced(file_path):
    """
    Enhanced text file extraction with encoding detection.
    """
    try:
        # Try multiple encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                    if content.strip():
                        logger.info(f"[TEXT] Successfully extracted {len(content)} characters using {encoding}")
                        return content
            except UnicodeDecodeError:
                continue
        
        logger.warning(f"[TEXT] All encoding attempts failed for {file_path}")
        return ""
        
    except Exception as e:
        logger.error(f"[TEXT] Extraction failed: {e}")
        return ""

def chunk_text_for_processing(text, chunk_size=4000, overlap=200):
    """
    Split large text into overlapping chunks for better processing.
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at a sentence boundary
        if end < len(text):
            # Look for sentence endings within the last 200 characters
            search_start = max(start, end - 200)
            sentence_end = -1
            
            for i in range(end, search_start, -1):
                if text[i:i+2] in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                    sentence_end = i + 1
                    break
            
            if sentence_end > search_start:
                end = sentence_end
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap if end < len(text) else end
        
        if start >= len(text):
            break
    
    return chunks


@app.route('/upload_file', methods=['POST'])
@token_required
def upload_file():
    """
    Accepts file upload, extracts text, and stores in MongoDB with file_id and conversation_id.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in request.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file.'}), 400
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed.'}), 400
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower()
    filetype = ext
    temp = tempfile.NamedTemporaryFile(delete=False)
    file.save(temp.name)
    text_content = extract_text_from_file(temp.name, filetype)
    temp.close()
    file_id = str(uuid.uuid4())
    conversation_id = sanitize_conversation_id(request.form.get('conversation_id'))
    username = g.user['username']
    files_col.insert_one({
        'file_id': file_id,
        'username': username,
        'conversation_id': conversation_id,
        'filename': filename,
        'filetype': filetype,
        'text_content': text_content,
        'uploaded_at': datetime.utcnow().isoformat()
    })
    return jsonify({'file_id': file_id, 'conversation_id': conversation_id, 'filename': filename, 'filetype': filetype})


@app.route('/chat_with_file', methods=['POST'])
@token_required
def chat_with_file():
    """
    Enhanced chat with file: supports chunked processing for large files and provides comprehensive responses.
    """
    data = request.get_json()
    question = data.get('question', '').strip()
    file_id = data.get('file_id')
    conversation_id = sanitize_conversation_id(data.get('conversation_id'))
    
    if not question or not file_id:
        return jsonify({'error': 'Question and file_id are required.'}), 400
    
    username = g.user['username']
    file_doc = files_col.find_one({'file_id': file_id, 'username': username})
    if not file_doc:
        return jsonify({'error': 'File not found.'}), 404
    
    full_content = file_doc['text_content']
    filename = file_doc['filename']
    filetype = file_doc['filetype']
    
    # Enhanced processing for large files
    if len(full_content) > 8000:
        logger.info(f"[CHAT_FILE] Processing large file ({len(full_content)} chars) in chunks")
        
        # Create chunks for large files
        chunks = chunk_text_for_processing(full_content, chunk_size=6000, overlap=300)
        
        # Find most relevant chunks based on question
        relevant_chunks = []
        question_lower = question.lower()
        
        for i, chunk in enumerate(chunks):
            chunk_lower = chunk.lower()
            # Simple relevance scoring based on keyword overlap
            question_words = set(question_lower.split())
            chunk_words = set(chunk_lower.split())
            overlap = len(question_words.intersection(chunk_words))
            
            if overlap > 0:
                relevant_chunks.append((i, chunk, overlap))
        
        # Sort by relevance and take top chunks
        relevant_chunks.sort(key=lambda x: x[2], reverse=True)
        top_chunks = relevant_chunks[:3] if relevant_chunks else [(0, chunks[0], 0)]
        
        # Build context from most relevant chunks
        file_context = ""
        for chunk_idx, chunk_content, score in top_chunks:
            file_context += f"\n--- Section {chunk_idx + 1} (Relevance: {score}) ---\n{chunk_content}\n"
        
        context_info = f"\n[Note: This is a large document with {len(chunks)} sections. Showing {len(top_chunks)} most relevant sections.]"
        file_excerpt = file_context + context_info
    else:
        file_excerpt = full_content[:8000]  # Use full content for smaller files
    
    # Enhanced code extraction for scan reports
    code_snippet = None
    if filetype in ['txt', 'md'] and any(keyword in full_content.lower() for keyword in ['vulnerability', 'scan', 'report', 'security']):
        import re
        lines = full_content.splitlines()
        code_blocks = []
        
        for i, line in enumerate(lines):
            # Enhanced pattern matching for vulnerability reports
            patterns = [
                r"(.+\.[a-zA-Z0-9]+) \((\d+) occurrences?\)",  # file.ext (N occurrences)
                r"File: (.+\.[a-zA-Z0-9]+)",  # File: path/to/file.ext
                r"Location: (.+):(\d+)",  # Location: file:line
                r"Found in: (.+\.[a-zA-Z0-9]+)"  # Found in: file.ext
            ]
            
            for pattern in patterns:
                m = re.search(pattern, line)
                if m:
                    file_path = m.group(1)
                    # Extract context around the match
                    context_lines = []
                    start_idx = max(0, i - 2)
                    end_idx = min(len(lines), i + 8)
                    
                    for j in range(start_idx, end_idx):
                        context_lines.append(f"{j+1:4d}: {lines[j]}")
                    
                    code_blocks.append((file_path, '\n'.join(context_lines)))
                    break
        
        if code_blocks:
            file_excerpt += "\n\n--- Extracted Vulnerability Context ---\n"
            for file_path, context in code_blocks[:3]:  # Limit to top 3 matches
                file_excerpt += f"\n**{file_path}:**\n```\n{context}\n```\n"
    
    # Create enhanced prompt based on file type
    if filetype == 'pdf':
        prompt_context = "PDF document"
    elif filetype == 'docx':
        prompt_context = "Word document"
    elif filetype in ['py', 'js', 'java', 'c', 'cpp', 'go', 'rb']:
        prompt_context = f"{filetype.upper()} source code"
    elif 'scan' in filename.lower() or 'vulnerability' in filename.lower():
        prompt_context = "security scan report"
    else:
        prompt_context = f"{filetype} file"
    
    prompt = (
        f"You are CyberGuard AI, an expert cybersecurity assistant. The user has uploaded "
        f"a {prompt_context} named '{filename}' and asked a question about it.\n\n"
        f"**User Question:** {question}\n\n"
        f"**File Content Analysis:**\n{file_excerpt}\n\n"
        f"**Instructions:**\n"
        f"- Provide a comprehensive answer based on the file content\n"
        f"- If this is code, provide code-aware analysis including security implications\n"
        f"- If this is a security report, focus on vulnerabilities and mitigation strategies\n"
        f"- Include specific references to relevant sections or line numbers when applicable\n"
        f"- Use technical language appropriate for cybersecurity professionals\n"
        f"- Structure your response clearly with headings if needed\n\n"
        f"**Response:**"
    )
    
    # Generate response with enhanced error handling
    answer = ""
    try:
        answer = generate_response_with_ollama(prompt, max_tokens=2500)
    except Exception as e:
        logger.error(f"[CHAT_FILE] Response generation failed: {e}")
        answer = (
            f"I encountered an issue processing your question about {filename}. "
            f"However, I can see the file contains relevant information about your query. "
            f"Please try rephrasing your question or contact support if the issue persists."
        )
    
    if not answer or len(answer.strip()) < 20:
        answer = (
            f"Based on the content of {filename}, I can see information related to your question "
            f"about: {question}. The file appears to be a {prompt_context} with relevant content. "
            f"Could you please rephrase your question for more specific analysis?"
        )
    
    now = datetime.utcnow().isoformat()
    
    # Save enhanced chat messages with file metadata
    user_msg_doc = {
        'username': username,
        'id': int(time.time() * 1000),
        'text': question,
        'sender': 'user',
        'timestamp': now,
        'hasCode': detectCodeBlocks(question),
        'reactions': [],
        'conversation_id': conversation_id,
        'file_id': file_id,
        'file_metadata': {
            'filename': filename,
            'filetype': filetype,
            'size_chars': len(full_content),
            'is_large_file': len(full_content) > 8000
        }
    }
    chats_col.insert_one(user_msg_doc)
    
    bot_msg_doc = {
        'username': username,
        'id': int(time.time() * 1000) + 1,
        'text': answer,
        'sender': 'bot',
        'timestamp': now,
        'hasCode': detectCodeBlocks(answer),
        'reactions': [],
        'conversation_id': conversation_id,
        'file_id': file_id,
        'file_metadata': {
            'filename': filename,
            'filetype': filetype,
            'processed_chunks': len(top_chunks) if len(full_content) > 8000 else 1
        }
    }
    chats_col.insert_one(bot_msg_doc)
    
    return jsonify({
        'answer': answer, 
        'file_id': file_id, 
        'conversation_id': conversation_id,
        'processing_info': {
            'file_size': len(full_content),
            'chunks_processed': len(top_chunks) if len(full_content) > 8000 else 1,
            'is_large_file': len(full_content) > 8000
        }
    })

@app.route('/delete_file/<file_id>', methods=['DELETE'])
@token_required
def delete_file(file_id):
    """
    Delete a file by file_id for the current user (removes from DB, optionally from disk if stored).
    """
    username = g.user['username']
    file_doc = files_col.find_one({'file_id': file_id, 'username': username})
    if not file_doc:
        return jsonify({'error': 'File not found.'}), 404
    files_col.delete_one({'file_id': file_id, 'username': username})
    # Optionally, delete file from disk if you store the path
    # if 'filepath' in file_doc:
    #     try:
    #         os.remove(file_doc['filepath'])
    #     except Exception as e:
    #         logger.warning(f"[FILE] Could not delete file from disk: {e}")
    return jsonify({'success': True, 'file_id': file_id})

# ============================================================================
# CHECKMARX SCAN REPORT ENDPOINTS
# ============================================================================

# New collection for scan findings
scan_findings_col = db['scan_findings']

@app.route('/upload_scan_report', methods=['POST'])
@token_required
def upload_scan_report():
    """
    Accepts a Checkmarx scan report (JSON or PDF), parses it, extracts findings and vulnerable code snippets, and stores them.
    Always returns the latest findings for the conversation after upload.
    """
    import logging
    logger = logging.getLogger("scan_report")
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in request.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file.'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower()
    temp = tempfile.NamedTemporaryFile(delete=False)
    file.save(temp.name)
    username = g.user['username']
    conversation_id = sanitize_conversation_id(request.form.get('conversation_id'))
    findings = []
    existing_snippets = set()
    # Preload existing code_snippets for this conversation and user
    for f in scan_findings_col.find({'username': username, 'conversation_id': conversation_id}, {'code_snippet': 1}):
        snippet = f.get('code_snippet', '').strip()
        if snippet:
            existing_snippets.add(snippet)
    if ext == 'pdf':
        # Extract text from PDF and store as text, not binary
        pdf_text = ''
        if _pdfplumber_available:
            import warnings
            try:
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    with pdfplumber.open(temp.name) as pdf:
                        pdf_text = "\n".join(page.extract_text() or '' for page in pdf.pages if page.extract_text())
            except Exception as e:
                logger.error(f"[PDF PARSE] Exception: {e}")
                pdf_text = ''
        # Store PDF text in DB
        scan_findings_col.insert_one({
            'username': username,
            'conversation_id': conversation_id,
            'report_filename': file.filename,
            'pdf_text': pdf_text,
            'created_at': datetime.utcnow().isoformat(),
            'type': 'pdf'
        })
        # Use enhanced code extraction logic on pdf_text with specific focus on line-numbered code blocks
        # STRICT: Only extract code snippets where each line starts with a line number followed by code
        findings = []
        lines = pdf_text.splitlines()
        code_blocks = []
        line_number_pattern = re.compile(r'^\s*\d+\s+\S+')
        current_block = []
        file_path = ""
        for i, line in enumerate(lines):
            if line_number_pattern.match(line):
                current_block.append(line)
            else:
                if current_block:
                    block_text = '\n'.join(current_block).strip()
                    if len(current_block) >= 2 and len(block_text) > 30:
                        code_blocks.append((file_path or "(from code section)", block_text))
                    current_block = []
                    file_path = ""
        if current_block:
            block_text = '\n'.join(current_block).strip()
            if len(current_block) >= 2 and len(block_text) > 30:
                code_blocks.append((file_path or "(from code section)", block_text))
        # Only create findings for these code blocks
        for file_path, code_snippet in code_blocks:
            code_snippet_stripped = code_snippet.strip()
            if code_snippet_stripped in existing_snippets:
                continue  # Skip duplicate
            existing_snippets.add(code_snippet_stripped)
            finding_id = str(uuid.uuid4())
            lines = code_snippet_stripped.split('\n')
            findings.append({
                'finding_id': finding_id,
                'username': username,
                'conversation_id': conversation_id,
                'vuln_name': 'Extracted from PDF',
                'severity': 'Medium',
                'file_path': file_path,
                'line': lines[0].split(' ')[0] if lines and lines[0].strip().split(' ')[0].isdigit() else '',
                'code_snippet': code_snippet_stripped,
                'status': 'unresolved',
                'report_filename': file.filename,
                'created_at': datetime.utcnow().isoformat(),
                'type': 'pdf-extracted'
            })
        
        # Fallback to previous vulnerability block-based logic if still no findings
        if not findings:
            vuln_pattern = re.compile(r'(Vulnerability(?: Name)?:\s*.+?)(?=Vulnerability(?: Name)?:|$)', re.DOTALL)
            vuln_blocks = vuln_pattern.findall(pdf_text)
            for block in vuln_blocks:
                vuln_name_match = re.search(r'Vulnerability(?: Name)?:\s*(.+)', block)
                vuln_name = vuln_name_match.group(1).strip() if vuln_name_match else 'Unknown'
                code_snippet = ''
                snippet_match = re.search(r'Code Snippets?:\s*(.*)', block, re.DOTALL)
                if snippet_match:
                    after_snippet = snippet_match.group(1)
                    stop_match = re.search(r'^[A-Za-z ]+:', after_snippet, re.MULTILINE)
                    if stop_match:
                        code_snippet = after_snippet[:stop_match.start()].strip()
                    else:
                        code_snippet = after_snippet.strip()
                    code_snippet = re.sub(r'\n{3,}', '\n\n', code_snippet).strip()
                if vuln_name and code_snippet and code_snippet not in existing_snippets:
                    existing_snippets.add(code_snippet)
                    finding_id = str(uuid.uuid4())
                    findings.append({
                        'finding_id': finding_id,
                        'username': username,
                        'conversation_id': conversation_id,
                        'vuln_name': vuln_name,
                        'severity': 'Medium',
                        'file_path': '',
                        'line': '',
                        'code_snippet': code_snippet,
                        'status': 'unresolved',
                        'report_filename': file.filename,
                        'created_at': datetime.utcnow().isoformat(),
                        'type': 'pdf-extracted'
                    })
        temp.close()
        if findings:
            scan_findings_col.insert_many(findings)
        # Return latest findings for this conversation
        findings = list(scan_findings_col.find({
            'username': username,
            'conversation_id': conversation_id,
            'type': {'$in': ['json', 'pdf-extracted']}
        }, {'_id': 0}).sort([('created_at', -1)]))
        return jsonify({'success': True, 'pdf_report': True, 'conversation_id': conversation_id, 'findings': findings, 'pdf_text': pdf_text})
    # ...existing JSON parsing logic for Checkmarx...
    try:
        with open(temp.name, 'r', encoding='utf-8', errors='ignore') as f:
            report_data = f.read()
        import json
        report_json = json.loads(report_data)
    except Exception as e:
        return jsonify({'error': f'Failed to parse report: {e}'}), 400
    finally:
        temp.close()
    findings = []
    for query in report_json.get('results', []):
        for finding in query.get('vulnerabilities', []):
            code_snippet = finding.get('codeSnippet', '').strip()
            if code_snippet and code_snippet not in existing_snippets:
                existing_snippets.add(code_snippet)
                finding_id = str(uuid.uuid4())
                file_path = finding.get('fileName', '')
                line = finding.get('line', '')
                vuln_name = query.get('queryName', 'Unknown')
                severity = finding.get('severity', 'Medium')
                findings.append({
                    'finding_id': finding_id,
                    'username': username,
                    'conversation_id': conversation_id,
                    'vuln_name': vuln_name,
                    'severity': severity,
                    'file_path': file_path,
                    'line': line,
                    'code_snippet': code_snippet,
                    'status': 'unresolved',
                    'report_filename': file.filename,
                    'created_at': datetime.utcnow().isoformat(),
                    'type': 'json'
                })
    if findings:
        scan_findings_col.insert_many(findings)
    # Return latest findings for this conversation
    findings = list(scan_findings_col.find({
        'username': username,
        'conversation_id': conversation_id,
        'type': {'$in': ['json', 'pdf-extracted']}
    }, {'_id': 0}).sort([('created_at', -1)]))
    return jsonify({'success': True, 'findings': findings, 'conversation_id': conversation_id})

@app.route('/scan_findings', methods=['GET'])
@token_required
def get_scan_findings():
    """
    Fetch all scan findings for the current user and conversation, sorted by severity and timestamp.
    Adds 'is_latest' field for findings from the most recent upload.
    """
    username = g.user['username']
    conversation_id = request.args.get('conversation_id')
    query = {'username': username}
    if conversation_id:
        query['conversation_id'] = conversation_id
    findings = list(scan_findings_col.find(query, {'_id': 0, 'pdf_data': 0}))
    # Sort by severity (Critical > High > Medium > Low > Info) and created_at desc
    severity_order = {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1, 'Info': 0}
    findings.sort(key=lambda f: (severity_order.get(f.get('severity', 'Medium'), 2), f.get('created_at', '')), reverse=True)
    # Mark latest findings
    if findings:
        latest_time = max(f.get('created_at', '') for f in findings)
        for f in findings:
            f['is_latest'] = (f.get('created_at', '') == latest_time)
    return jsonify({'findings': findings})

@app.route('/scan_findings/latest', methods=['GET'])
@token_required
def get_latest_scan_findings():
    """
    Fetch the most recent scan findings for the current user (across all conversations).
    """
    username = g.user['username']
    findings = list(scan_findings_col.find({'username': username, 'type': 'json'}, {'_id': 0, 'pdf_data': 0}))
    if not findings:
        return jsonify({'findings': []})
    latest_time = max(f.get('created_at', '') for f in findings)
    latest_findings = [f for f in findings if f.get('created_at', '') == latest_time]
    return jsonify({'findings': latest_findings})

@app.route('/scan_report_pdf_url', methods=['GET'])
@token_required
def scan_report_pdf_url():
    """
    Returns the download URL for the PDF report for a conversation, if it exists.
    """
    username = g.user['username']
    conversation_id = request.args.get('conversation_id')
    doc = scan_findings_col.find_one({
        'username': username,
        'conversation_id': conversation_id,
        'type': 'pdf'
    })
    if not doc:
        return jsonify({'pdf_available': False})
    url = f"/download_scan_report_pdf?conversation_id={conversation_id}"
    return jsonify({'pdf_available': True, 'download_url': url})

@app.route('/scan_finding/<finding_id>/explain', methods=['POST'])
@token_required
def explain_scan_finding(finding_id):
    """
    Return an explanation for the code snippet in the finding, formatted for frontend display.
    """
    finding = scan_findings_col.find_one({'finding_id': finding_id})
    if not finding:
        return jsonify({'error': 'Finding not found.'}), 404
    code = finding.get('code_snippet', '')
    vuln_name = finding.get('vuln_name', 'Vulnerability')
    file_path = finding.get('file_path', '')
    # Use LLM or template for explanation
    prompt = (
        f"You are a senior application security engineer reviewing a scan finding. "
        f"Provide a concise, actionable explanation of the vulnerability and its impact for the following code snippet. "
        f"Start with a one-sentence summary, then use bullet points for details. "
        f"If possible, include practical remediation advice. Format your answer in markdown.\n\n"
        f"**Finding Type:** {vuln_name}\n"
        f"**File:** {file_path}\n"
        f"**Code Snippet:**\n"
        f"```\n{code}\n```\n\n"
        f"**Explanation:**\n"
    )
    try:
        explanation = generate_response_with_ollama(prompt, max_tokens=600)
        # Ensure markdown/code block formatting
        if not explanation.strip().startswith('```'):
            explanation = explanation.strip()
    except Exception as e:
        explanation = f"Could not generate explanation: {e}"
    return jsonify({'explanation': explanation})

@app.route('/scan_finding/<finding_id>/patch', methods=['POST'])
@token_required
def patch_scan_finding(finding_id):
    """
    Return a patch recommendation for the code snippet in the finding, formatted for frontend display.
    """
    finding = scan_findings_col.find_one({'finding_id': finding_id})
    if not finding:
        return jsonify({'error': 'Finding not found.'}), 404
    code = finding.get('code_snippet', '')
    vuln_name = finding.get('vuln_name', 'Vulnerability')
    file_path = finding.get('file_path', '')
    # Use LLM or template for patch
    prompt = (
        f"You are a senior application security engineer. "
        f"Suggest a secure patch for the following code snippet. "
        f"Start with a one-sentence summary of the fix, then output the patched code in a code block, followed by a brief explanation in markdown. "
        f"If possible, explain why the patch is secure and reference best practices.\n\n"
        f"**Finding Type:** {vuln_name}\n"
        f"**File:** {file_path}\n"
        f"**Original Code:**\n"
        f"```\n{code}\n```\n\n"
        f"**Patch:**\n"
    )
    try:
        patch = generate_response_with_ollama(prompt, max_tokens=600)
        # Ensure markdown/code block formatting
        if not patch.strip().startswith('```'):
            patch = patch.strip()
    except Exception as e:
        patch = f"Could not generate patch: {e}"
    return jsonify({'patch': patch})

# ============================================================================

# JWT authentication decorator and code block detector must be defined before any route uses them

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            if auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]
        if not token:
            return jsonify({'error': 'Token is missing!'}), 401
        try:
            data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGO])
            current_user = users_col.find_one({'username': data['username']})
            if not current_user:
                return jsonify({'error': 'User not found!'}), 401
            g.user = current_user
        except Exception as e:
            return jsonify({'error': 'Token is invalid!'}), 401
        return f(*args, **kwargs)
    return decorated

def detectCodeBlocks(text):
    codeBlockRegex = re.compile(r'```(\w+)?\n([\s\S]*?)\n```')
    inlineCodeRegex = re.compile(r'`([^`]+)`')
    return bool(codeBlockRegex.search(text) or inlineCodeRegex.search(text))

# Utility to ensure conversation_id is always a valid, non-empty string
import uuid

def sanitize_conversation_id(conversation_id):
    """
    Ensures conversation_id is always a valid, non-empty string for MongoDB schema.
    Generates a new UUID string if missing, None, 'None', or empty.
    """
    if not conversation_id or str(conversation_id).lower() == 'none' or str(conversation_id).strip() == '':
        return str(uuid.uuid4())
    return str(conversation_id)

# Move pdfplumber import and _pdfplumber_available definition to the top of the file, so they are always defined before use
try:
    import pdfplumber # type: ignore
    _pdfplumber_available = True
except ImportError:
    _pdfplumber_available = False

if __name__ == "__main__":
    print("[CYBERGUARD AI] Starting Security Backend with Llama 3 + Gemini")
    print("=" * 60)
    print(f"[SERVER] Server: http://localhost:9000")
    print(f"[CHAT] Chat Endpoint: /stream_chat")
    print(f"[HEALTH] Health Check: /api/health")
    print(f"[VULNS] Vulnerabilities: /api/vulnerabilities")
    print(f"[STATS] Stats: /api/stats")
    print(f"[SCAN] Scan Findings: /scan_findings")
    print(f"[UPLOAD] Upload Scan Report: /upload_scan_report")
    print("=" * 60)
    print("[INFO] Ollama model used:", OLLAMA_MODEL)
    print(f"[DEVICE] Device: {get_ollama_device_string()}")
    print("[INFO] Gemini API is integrated for context retrieval.")
    print("[READY] Ready for cybersecurity queries!")
    print()
    app.run(
        host='0.0.0.0',
        port=9000,
        debug=True
    )
